"""Document service for handling file uploads and text extraction"""
import os
import uuid
from datetime import datetime
from typing import Dict, Any, Optional, List
from pathlib import Path
import logging

# Document processing libraries
from pypdf import PdfReader
from docx import Document

# Import chunking and vector store services
from app.services.embedding_service import embedding_service
from app.services.vector_store_service import vector_store_service

# Import error handlers
from app.utils.error_handlers import (
    DocumentError,
    DocumentNotFoundError,
    InvalidFileFormatError,
    EmptyDocumentError,
    validate_file_format,
    validate_non_empty_text,
    log_error
)

logger = logging.getLogger(__name__)


class DocumentMetadata:
    """Metadata for uploaded documents"""
    
    def __init__(
        self,
        file_id: str,
        original_filename: str,
        stored_filename: str,
        file_type: str,
        upload_time: str,
        extracted_text_length: int,
        file_size: int
    ):
        self.file_id = file_id
        self.original_filename = original_filename
        self.stored_filename = stored_filename
        self.file_type = file_type
        self.upload_time = upload_time
        self.extracted_text_length = extracted_text_length
        self.file_size = file_size
    
    def to_dict(self) -> Dict[str, Any]:
        """Convert metadata to dictionary"""
        return {
            "file_id": self.file_id,
            "original_filename": self.original_filename,
            "stored_filename": self.stored_filename,
            "file_type": self.file_type,
            "upload_time": self.upload_time,
            "extracted_text_length": self.extracted_text_length,
            "file_size": self.file_size
        }


class DocumentService:
    """Service for handling document uploads and text extraction"""
    
    # Supported file extensions
    SUPPORTED_EXTENSIONS = {".pdf", ".docx"}
    
    def __init__(self, upload_dir: str = "uploads"):
        """
        Initialize document service
        
        Args:
            upload_dir: Directory to store uploaded files
        """
        self.upload_dir = Path(upload_dir)
        self.upload_dir.mkdir(exist_ok=True)
        
        # In-memory storage for document metadata and text
        # In production, use a database
        self.documents: Dict[str, Dict[str, Any]] = {}
        
        logger.info(f"DocumentService initialized with upload directory: {self.upload_dir}")
    
    def validate_file_extension(self, filename: str) -> None:
        """
        Validate if file extension is supported
        
        Args:
            filename: Name of the file
            
        Raises:
            InvalidFileFormatError: If file format is not supported
        """
        file_ext = Path(filename).suffix.lower()
        
        if not file_ext:
            raise InvalidFileFormatError(
                file_format="no extension",
                supported_formats=list(self.SUPPORTED_EXTENSIONS)
            )
        
        if file_ext not in self.SUPPORTED_EXTENSIONS:
            raise InvalidFileFormatError(
                file_format=file_ext,
                supported_formats=list(self.SUPPORTED_EXTENSIONS)
            )
    
    def generate_unique_filename(self, original_filename: str) -> str:
        """
        Generate a unique filename to avoid collisions
        
        Args:
            original_filename: Original name of the uploaded file
            
        Returns:
            Unique filename
        """
        file_ext = Path(original_filename).suffix.lower()
        unique_id = uuid.uuid4().hex[:12]
        timestamp = datetime.utcnow().strftime("%Y%m%d_%H%M%S")
        
        return f"{timestamp}_{unique_id}{file_ext}"
    
    async def extract_text_from_pdf(self, file_path: Path) -> str:
        """
        Extract text from PDF file
        
        Args:
            file_path: Path to PDF file
            
        Returns:
            Extracted text
        """
        try:
            reader = PdfReader(str(file_path))
            text_parts = []
            
            for page_num, page in enumerate(reader.pages, 1):
                try:
                    text = page.extract_text()
                    if text:
                        text_parts.append(text)
                except Exception as e:
                    logger.warning(f"Error extracting text from page {page_num}: {str(e)}")
                    continue
            
            full_text = "\n".join(text_parts)
            logger.info(f"Extracted {len(full_text)} characters from PDF with {len(reader.pages)} pages")
            
            return full_text
            
        except Exception as e:
            logger.error(f"Error reading PDF: {str(e)}")
            raise Exception(f"Failed to extract text from PDF: {str(e)}")
    
    async def extract_text_from_docx(self, file_path: Path) -> str:
        """
        Extract text from DOCX file
        
        Args:
            file_path: Path to DOCX file
            
        Returns:
            Extracted text
        """
        try:
            doc = Document(str(file_path))
            text_parts = []
            
            # Extract text from paragraphs
            for para in doc.paragraphs:
                if para.text.strip():
                    text_parts.append(para.text)
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            text_parts.append(cell.text)
            
            full_text = "\n".join(text_parts)
            logger.info(f"Extracted {len(full_text)} characters from DOCX")
            
            return full_text
            
        except Exception as e:
            logger.error(f"Error reading DOCX: {str(e)}")
            raise Exception(f"Failed to extract text from DOCX: {str(e)}")
    
    async def save_and_process_file(
        self,
        file_content: bytes,
        filename: str
    ) -> Dict[str, Any]:
        """
        Save uploaded file and extract text
        
        Args:
            file_content: Binary content of the file
            filename: Original filename
            
        Returns:
            Dictionary with file metadata and extracted text
        """
        # Validate file extension
        is_valid, error_msg = self.validate_file_extension(filename)
        if not is_valid:
            raise ValueError(error_msg)
        
        # Generate unique filename
        stored_filename = self.generate_unique_filename(filename)
        file_path = self.upload_dir / stored_filename
        
        # Save file to disk
        try:
            with open(file_path, "wb") as f:
                f.write(file_content)
            logger.info(f"File saved: {stored_filename}")
        except Exception as e:
            logger.error(f"Error saving file: {str(e)}")
            raise Exception(f"Failed to save file: {str(e)}")
        
        # Extract text based on file type
        file_ext = Path(filename).suffix.lower()
        try:
            if file_ext == ".pdf":
                extracted_text = await self.extract_text_from_pdf(file_path)
            elif file_ext == ".docx":
                extracted_text = await self.extract_text_from_docx(file_path)
            else:
                raise ValueError(f"Unsupported file type: {file_ext}")
        except Exception as e:
            # Clean up file if text extraction fails
            if file_path.exists():
                file_path.unlink()
            raise e
        
        # Generate file ID
        file_id = uuid.uuid4().hex
        upload_time = datetime.utcnow().isoformat()
        
        # Create metadata
        metadata = DocumentMetadata(
            file_id=file_id,
            original_filename=filename,
            stored_filename=stored_filename,
            file_type=file_ext,
            upload_time=upload_time,
            extracted_text_length=len(extracted_text),
            file_size=len(file_content)
        )
        
        # Store document data
        self.documents[file_id] = {
            "metadata": metadata.to_dict(),
            "extracted_text": extracted_text
        }
        
        # Create chunks and store in vector database
        chunks = embedding_service.split_text_into_chunks(
            text=extracted_text,
            document_id=file_id,
            source_filename=filename
        )
        
        # Store chunks in vector database
        vector_result = vector_store_service.add_chunks(chunks)
        
        logger.info(f"Document processed successfully: {file_id}")
        logger.info(f"Created {len(chunks)} chunks and stored in vector database")
        
        return {
            "file_id": file_id,
            "metadata": metadata.to_dict(),
            "extracted_text": extracted_text,
            "chunks_created": len(chunks),
            "vector_store_result": vector_result
        }
    
    def get_document(self, file_id: str) -> Optional[Dict[str, Any]]:
        """
        Get document by ID
        
        Args:
            file_id: Document ID
            
        Returns:
            Document data or None if not found
        """
        return self.documents.get(file_id)
    
    def get_all_documents(self) -> List[Dict[str, Any]]:
        """
        Get all document metadata
        
        Returns:
            List of document metadata
        """
        return [
            doc["metadata"]
            for doc in self.documents.values()
        ]
    
    def delete_document(self, file_id: str) -> bool:
        """
        Delete a document by ID
        
        Args:
            file_id: Document ID
            
        Returns:
            True if deleted, False if not found
        """
        doc = self.documents.get(file_id)
        if not doc:
            return False
        
        # Delete physical file
        stored_filename = doc["metadata"]["stored_filename"]
        file_path = self.upload_dir / stored_filename
        
        if file_path.exists():
            try:
                file_path.unlink()
                logger.info(f"Deleted file: {stored_filename}")
            except Exception as e:
                logger.error(f"Error deleting file: {str(e)}")
        
        # Delete chunks from vector store
        vector_delete_result = vector_store_service.delete_document_chunks(file_id)
        logger.info(f"Deleted {vector_delete_result.get('chunks_deleted', 0)} chunks from vector store")
        
        # Remove from storage
        del self.documents[file_id]
        
        return True


# Singleton instance
document_service = DocumentService()
